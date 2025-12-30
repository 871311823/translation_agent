#!/usr/bin/env python3
"""
Translation Agent Pro - 桌面版
Desktop GUI Application for Batch Translation
"""

import os
import sys
import json
import threading
import time
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor, as_completed
from typing import List, Dict, Tuple, Optional
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
from tkinter.ttk import Progressbar, Notebook

# 添加 app 目录到路径
app_dir = os.path.join(os.path.dirname(__file__), 'app')
if app_dir not in sys.path:
    sys.path.insert(0, app_dir)

try:
    from process import (
        extract_docx, extract_pdf, extract_text,
        model_load, translator, translator_sec
    )
except ImportError as e:
    print(f"导入模块失败: {e}")
    print("请确保 app 目录下的相关文件存在")

# 配置文件路径
CONFIG_FILE = "translation_config.json"
MAX_CONCURRENT_TASKS = 5

class TranslationTask:
    """翻译任务类"""
    def __init__(self, task_id: str, filename: str, content: str, file_path: str):
        self.task_id = task_id
        self.filename = filename
        self.content = content
        self.file_path = file_path
        self.status = "等待中"
        self.progress = 0
        self.init_translation = ""
        self.reflect_translation = ""
        self.final_translation = ""
        self.error_message = ""
        self.start_time = None
        self.end_time = None


class TranslationAgentGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("Translation Agent Pro - 专业批量翻译软件")
        self.root.geometry("1200x800")
        self.root.minsize(1000, 700)
        
        # 设置图标和样式
        self.setup_styles()
        
        # 初始化变量
        self.translation_tasks = {}
        self.task_counter = 0
        self.scanned_files = []
        self.is_translating = False
        self.is_paused = False  # 暂停标志
        self.is_loading_config = False  # 标志：是否正在加载配置（防止 endpoint.change 覆盖模型名）
        
        # 创建界面
        self.create_widgets()
        
        # 加载配置
        self.load_config()
        
        # 设置关闭事件
        self.root.protocol("WM_DELETE_WINDOW", self.on_closing)
    
    def setup_styles(self):
        """设置界面样式"""
        style = ttk.Style()
        style.theme_use('clam')
        
        # 配置样式
        style.configure('Title.TLabel', font=('Arial', 16, 'bold'))
        style.configure('Heading.TLabel', font=('Arial', 12, 'bold'))
        style.configure('Success.TLabel', foreground='green')
        style.configure('Error.TLabel', foreground='red')
        style.configure('Warning.TLabel', foreground='orange')
        
        # 配置按钮样式
        style.configure('Accent.TButton', font=('Arial', 10, 'bold'))
        
        # 配置复选框样式
        style.configure('Switch.TCheckbutton', font=('Arial', 10))
        
        # 配置LabelFrame样式
        style.configure('TLabelframe', borderwidth=2, relief='groove')
        style.configure('TLabelframe.Label', font=('Arial', 11, 'bold'))
    
    def create_widgets(self):
        """创建主界面组件"""
        # 主标题
        title_frame = ttk.Frame(self.root)
        title_frame.pack(fill='x', padx=10, pady=5)
        
        ttk.Label(title_frame, text="🚀 Translation Agent Pro", 
                 style='Title.TLabel').pack()
        ttk.Label(title_frame, text="专业批量翻译软件 v2.0.0").pack()
        
        # 创建标签页
        self.notebook = Notebook(self.root)
        self.notebook.pack(fill='both', expand=True, padx=10, pady=5)
        
        # 创建各个页面
        self.create_api_settings_tab()
        self.create_file_management_tab()
        self.create_progress_tab()
        self.create_about_tab()
    
    def create_api_settings_tab(self):
        """创建API设置页面"""
        api_frame = ttk.Frame(self.notebook)
        self.notebook.add(api_frame, text="🔧 API设置")
        
        # 创建主容器，使用两列布局
        main_container = ttk.Frame(api_frame)
        main_container.pack(fill='both', expand=True, padx=15, pady=15)
        
        # 左侧配置区域 - 固定宽度
        left_frame = ttk.Frame(main_container)
        left_frame.pack(side='left', fill='y', padx=(0, 15))
        left_frame.configure(width=600)
        left_frame.pack_propagate(False)
        
        # 右侧状态和帮助区域 - 填充剩余空间
        right_frame = ttk.Frame(main_container)
        right_frame.pack(side='right', fill='both', expand=True, padx=(15, 0))
        
        # 创建滚动框架（左侧）
        canvas = tk.Canvas(left_frame, highlightthickness=0)
        scrollbar = ttk.Scrollbar(left_frame, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        # === 主要API配置区域 ===
        main_api_frame = ttk.LabelFrame(scrollable_frame, text="🌐 主要API配置", padding=20)
        main_api_frame.pack(fill='x', pady=(0, 15))
        
        # API端点选择行
        endpoint_row = ttk.Frame(main_api_frame)
        endpoint_row.pack(fill='x', pady=(0, 12))
        
        ttk.Label(endpoint_row, text="API端点:", font=('Arial', 10, 'bold')).pack(side='left')
        self.endpoint_var = tk.StringVar(value="OpenAI")
        endpoint_combo = ttk.Combobox(endpoint_row, textvariable=self.endpoint_var,
                                     values=["OpenAI", "Groq", "TogetherAI", "Ollama", "CUSTOM"],
                                     state="readonly", width=15, font=('Arial', 10))
        endpoint_combo.pack(side='left', padx=(10, 0))
        endpoint_combo.bind('<<ComboboxSelected>>', self.on_endpoint_change)
        
        # 端点状态指示器
        self.endpoint_status = ttk.Label(endpoint_row, text="⚪", font=('Arial', 12))
        self.endpoint_status.pack(side='right')
        
        # 模型名称行
        model_row = ttk.Frame(main_api_frame)
        model_row.pack(fill='x', pady=(0, 12))
        
        ttk.Label(model_row, text="模型名称:", font=('Arial', 10, 'bold')).pack(side='left')
        self.model_var = tk.StringVar(value="gpt-4o")
        model_entry = ttk.Entry(model_row, textvariable=self.model_var, width=25, font=('Arial', 10))
        model_entry.pack(side='left', padx=(10, 0), fill='x', expand=True)
        
        # 模型建议按钮
        ttk.Button(model_row, text="💡", width=3, 
                  command=self.show_model_suggestions).pack(side='right', padx=(5, 0))
        
        # API密钥行
        key_row = ttk.Frame(main_api_frame)
        key_row.pack(fill='x', pady=(0, 12))
        
        ttk.Label(key_row, text="API密钥:", font=('Arial', 10, 'bold')).pack(side='left')
        self.api_key_var = tk.StringVar()
        key_entry = ttk.Entry(key_row, textvariable=self.api_key_var, show="*", 
                             width=30, font=('Arial', 10))
        key_entry.pack(side='left', padx=(10, 0), fill='x', expand=True)
        
        # 显示/隐藏密钥按钮
        self.show_key_var = tk.BooleanVar()
        ttk.Checkbutton(key_row, text="👁️", variable=self.show_key_var,
                       command=lambda: key_entry.config(show="" if self.show_key_var.get() else "*")).pack(side='right', padx=(5, 0))
        
        # 基础URL行（条件显示）
        self.base_url_row = ttk.Frame(main_api_frame)
        
        ttk.Label(self.base_url_row, text="基础URL:", font=('Arial', 10, 'bold')).pack(side='left')
        self.base_url_var = tk.StringVar()
        ttk.Entry(self.base_url_row, textvariable=self.base_url_var, width=40, 
                 font=('Arial', 10)).pack(side='left', padx=(10, 0), fill='x', expand=True)
        
        # 连接测试按钮区域
        test_frame = ttk.Frame(main_api_frame)
        test_frame.pack(fill='x', pady=(15, 0))
        
        test_btn = ttk.Button(test_frame, text="🔍 测试API连接", 
                             command=self.test_api_connection, style='Accent.TButton')
        test_btn.pack(side='left')
        
        save_btn = ttk.Button(test_frame, text="💾 保存配置", 
                             command=self.save_config)
        save_btn.pack(side='left', padx=(10, 0))
        
        # 状态显示
        self.api_status_var = tk.StringVar(value="请配置API设置")
        self.api_status_label = ttk.Label(test_frame, textvariable=self.api_status_var,
                                         font=('Arial', 9))
        self.api_status_label.pack(side='right')
        
        # === 额外端点配置区域 ===
        extra_api_frame = ttk.LabelFrame(scrollable_frame, text="⚡ 额外端点配置（可选）", padding=20)
        extra_api_frame.pack(fill='x', pady=(0, 15))
        
        # 启用额外端点复选框
        extra_enable_frame = ttk.Frame(extra_api_frame)
        extra_enable_frame.pack(fill='x', pady=(0, 12))
        
        self.use_extra_endpoint_var = tk.BooleanVar()
        ttk.Checkbutton(extra_enable_frame, text="启用额外端点（用于反思步骤）",
                       variable=self.use_extra_endpoint_var,
                       command=self.toggle_extra_endpoint,
                       style='Switch.TCheckbutton').pack(side='left')
        
        # 额外端点详细配置
        self.extra_endpoint_frame = ttk.Frame(extra_api_frame)
        
        # 额外端点选择
        extra_endpoint_row = ttk.Frame(self.extra_endpoint_frame)
        extra_endpoint_row.pack(fill='x', pady=(0, 10))
        
        ttk.Label(extra_endpoint_row, text="额外端点:", font=('Arial', 10, 'bold')).pack(side='left')
        self.endpoint2_var = tk.StringVar(value="OpenAI")
        endpoint2_combo = ttk.Combobox(extra_endpoint_row, textvariable=self.endpoint2_var,
                    values=["OpenAI", "Groq", "TogetherAI", "Ollama", "CUSTOM"],
                    state="readonly", width=15, font=('Arial', 10))
        endpoint2_combo.pack(side='left', padx=(10, 0))
        endpoint2_combo.bind('<<ComboboxSelected>>', self.on_endpoint2_change)
        
        # 额外模型
        extra_model_row = ttk.Frame(self.extra_endpoint_frame)
        extra_model_row.pack(fill='x', pady=(0, 10))
        
        ttk.Label(extra_model_row, text="额外模型:", font=('Arial', 10, 'bold')).pack(side='left')
        self.model2_var = tk.StringVar(value="gpt-4o")
        ttk.Entry(extra_model_row, textvariable=self.model2_var, width=25, 
                 font=('Arial', 10)).pack(side='left', padx=(10, 0), fill='x', expand=True)
        
        # 额外密钥
        extra_key_row = ttk.Frame(self.extra_endpoint_frame)
        extra_key_row.pack(fill='x', pady=(0, 10))
        
        ttk.Label(extra_key_row, text="额外密钥:", font=('Arial', 10, 'bold')).pack(side='left')
        self.api_key2_var = tk.StringVar()
        extra_key_entry = ttk.Entry(extra_key_row, textvariable=self.api_key2_var, show="*", 
                                   width=30, font=('Arial', 10))
        extra_key_entry.pack(side='left', padx=(10, 0), fill='x', expand=True)
        
        # 额外密钥显示/隐藏
        self.show_key2_var = tk.BooleanVar()
        ttk.Checkbutton(extra_key_row, text="👁️", variable=self.show_key2_var,
                       command=lambda: extra_key_entry.config(show="" if self.show_key2_var.get() else "*")).pack(side='right', padx=(5, 0))
        
        # 额外基础URL行（条件显示）
        self.base_url2_row = ttk.Frame(self.extra_endpoint_frame)
        
        ttk.Label(self.base_url2_row, text="额外基础URL:", font=('Arial', 10, 'bold')).pack(side='left')
        self.base_url2_var = tk.StringVar()
        ttk.Entry(self.base_url2_row, textvariable=self.base_url2_var, width=40, 
                 font=('Arial', 10)).pack(side='left', padx=(10, 0), fill='x', expand=True)
        
        # === 翻译参数配置区域 ===
        translation_frame = ttk.LabelFrame(scrollable_frame, text="🌍 翻译参数", padding=20)
        translation_frame.pack(fill='x', pady=(0, 15))
        
        # 语言配置行
        lang_row = ttk.Frame(translation_frame)
        lang_row.pack(fill='x', pady=(0, 12))
        
        # 源语言
        ttk.Label(lang_row, text="源语言:", font=('Arial', 10, 'bold')).pack(side='left')
        self.source_lang_var = tk.StringVar(value="Chinese")
        ttk.Entry(lang_row, textvariable=self.source_lang_var, width=12, 
                 font=('Arial', 10)).pack(side='left', padx=(10, 20))
        
        # 目标语言
        ttk.Label(lang_row, text="目标语言:", font=('Arial', 10, 'bold')).pack(side='left')
        self.target_lang_var = tk.StringVar(value="English")
        ttk.Entry(lang_row, textvariable=self.target_lang_var, width=12, 
                 font=('Arial', 10)).pack(side='left', padx=(10, 20))
        
        # 地区
        ttk.Label(lang_row, text="地区:", font=('Arial', 10, 'bold')).pack(side='left')
        self.country_var = tk.StringVar(value="United States")
        ttk.Entry(lang_row, textvariable=self.country_var, width=15, 
                 font=('Arial', 10)).pack(side='left', padx=(10, 0))
        
        # === 高级参数配置区域 ===
        advanced_frame = ttk.LabelFrame(scrollable_frame, text="⚙️ 高级参数", padding=20)
        advanced_frame.pack(fill='x', pady=(0, 15))
        
        # Token数配置
        token_frame = ttk.Frame(advanced_frame)
        token_frame.pack(fill='x', pady=(0, 10))
        
        ttk.Label(token_frame, text="最大Token数:", font=('Arial', 10, 'bold')).pack(side='left')
        self.max_tokens_var = tk.IntVar(value=1000)
        
        # Token数输入框
        token_spinbox = ttk.Spinbox(token_frame, from_=512, to=4096, 
                                   textvariable=self.max_tokens_var, 
                                   width=6, font=('Arial', 10))
        token_spinbox.pack(side='left', padx=(10, 10))
        
        # Token数滑动条
        token_scale = ttk.Scale(token_frame, from_=512, to=4096, variable=self.max_tokens_var,
                               orient='horizontal', length=200, command=self.on_token_change)
        token_scale.pack(side='left', padx=(0, 10))
        
        # Token数标签
        self.token_label = ttk.Label(token_frame, textvariable=self.max_tokens_var, 
                                    font=('Arial', 10, 'bold'), foreground='blue')
        self.token_label.pack(side='left')
        
        # 温度配置
        temp_frame = ttk.Frame(advanced_frame)
        temp_frame.pack(fill='x', pady=(0, 10))
        
        ttk.Label(temp_frame, text="温度:", font=('Arial', 10, 'bold')).pack(side='left')
        self.temperature_var = tk.DoubleVar(value=0.3)
        
        # 温度输入框
        temp_spinbox = ttk.Spinbox(temp_frame, from_=0.0, to=1.0, increment=0.1,
                                  textvariable=self.temperature_var, 
                                  width=6, font=('Arial', 10), format="%.1f")
        temp_spinbox.pack(side='left', padx=(10, 10))
        
        # 温度滑动条
        temp_scale = ttk.Scale(temp_frame, from_=0.0, to=1.0, variable=self.temperature_var,
                              orient='horizontal', length=200, command=self.on_temp_change)
        temp_scale.pack(side='left', padx=(0, 10))
        
        # 温度标签
        self.temp_label = ttk.Label(temp_frame, text="0.3", 
                                   font=('Arial', 10, 'bold'), foreground='green')
        self.temp_label.pack(side='left')
        
        # RPM配置
        rpm_frame = ttk.Frame(advanced_frame)
        rpm_frame.pack(fill='x', pady=(0, 10))
        
        ttk.Label(rpm_frame, text="每分钟请求数:", font=('Arial', 10, 'bold')).pack(side='left')
        self.rpm_var = tk.IntVar(value=60)
        
        # RPM输入框
        rpm_spinbox = ttk.Spinbox(rpm_frame, from_=1, to=1000, 
                                 textvariable=self.rpm_var, 
                                 width=6, font=('Arial', 10))
        rpm_spinbox.pack(side='left', padx=(10, 10))
        
        # RPM滑动条
        rpm_scale = ttk.Scale(rpm_frame, from_=1, to=1000, variable=self.rpm_var,
                             orient='horizontal', length=200, command=self.on_rpm_change)
        rpm_scale.pack(side='left', padx=(0, 10))
        
        # RPM标签
        self.rpm_label = ttk.Label(rpm_frame, textvariable=self.rpm_var, 
                                  font=('Arial', 10, 'bold'), foreground='red')
        self.rpm_label.pack(side='left')
        
        # === 性能优化配置区域 ===
        performance_frame = ttk.LabelFrame(scrollable_frame, text="🚀 性能优化", padding=20)
        performance_frame.pack(fill='x', pady=(0, 15))
        
        # 超时设置
        timeout_frame = ttk.Frame(performance_frame)
        timeout_frame.pack(fill='x', pady=(0, 10))
        
        ttk.Label(timeout_frame, text="API超时(秒):", font=('Arial', 10, 'bold')).pack(side='left')
        self.api_timeout_var = tk.IntVar(value=300)  # 默认5分钟
        
        # 超时输入框
        timeout_spinbox = ttk.Spinbox(timeout_frame, from_=60, to=1800, 
                                     textvariable=self.api_timeout_var, 
                                     width=6, font=('Arial', 10))
        timeout_spinbox.pack(side='left', padx=(10, 10))
        
        # 超时滑动条
        timeout_scale = ttk.Scale(timeout_frame, from_=60, to=1800, variable=self.api_timeout_var,
                                 orient='horizontal', length=200, command=self.on_timeout_change)
        timeout_scale.pack(side='left', padx=(0, 10))
        
        # 超时标签
        self.timeout_label = ttk.Label(timeout_frame, textvariable=self.api_timeout_var, 
                                      font=('Arial', 10, 'bold'), foreground='orange')
        self.timeout_label.pack(side='left')
        
        # 超时建议
        ttk.Label(timeout_frame, text="(建议: 小文件180s, 大文件600s)", 
                 font=('Arial', 8), foreground='gray').pack(side='left', padx=(10, 0))
        
        # 性能模式选择
        mode_frame = ttk.Frame(performance_frame)
        mode_frame.pack(fill='x', pady=(0, 10))
        
        ttk.Label(mode_frame, text="性能模式:", font=('Arial', 10, 'bold')).pack(side='left')
        self.performance_mode_var = tk.StringVar(value="平衡")
        mode_combo = ttk.Combobox(mode_frame, textvariable=self.performance_mode_var,
                                 values=["快速", "平衡", "稳定"], state="readonly", width=10)
        mode_combo.pack(side='left', padx=(10, 20))
        mode_combo.bind('<<ComboboxSelected>>', self.on_performance_mode_change)
        
        # 模式说明
        self.mode_desc_label = ttk.Label(mode_frame, text="• 平衡: 适合大多数情况", 
                                        font=('Arial', 8), foreground='blue')
        self.mode_desc_label.pack(side='left', padx=(10, 0))
        
        # 重试设置
        retry_frame = ttk.Frame(performance_frame)
        retry_frame.pack(fill='x', pady=(0, 10))
        
        ttk.Label(retry_frame, text="失败重试次数:", font=('Arial', 10, 'bold')).pack(side='left')
        self.retry_count_var = tk.IntVar(value=2)
        
        retry_spinbox = ttk.Spinbox(retry_frame, from_=0, to=5, 
                                   textvariable=self.retry_count_var, 
                                   width=6, font=('Arial', 10))
        retry_spinbox.pack(side='left', padx=(10, 10))
        
        ttk.Label(retry_frame, text="次 (0=不重试)", 
                 font=('Arial', 8), foreground='gray').pack(side='left', padx=(10, 0))
        
        # 打包滚动区域（左侧）
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        # 右侧状态和帮助区域
        self.create_api_status_panel(right_frame)
    
    def create_api_status_panel(self, parent):
        """创建API状态和帮助面板"""
        # 不设置固定宽度，让它自动扩展
        # parent.configure(width=320)
        # parent.pack_propagate(False)
        
        # === API状态面板 ===
        status_frame = ttk.LabelFrame(parent, text="🔍 连接状态", padding=15)
        status_frame.pack(fill='both', expand=True, pady=(0, 15))
        
        # 状态指示器容器
        indicator_frame = ttk.Frame(status_frame)
        indicator_frame.pack(fill='x', pady=(0, 10))
        
        # 大型状态指示器
        self.status_indicator = ttk.Label(indicator_frame, text="⚪", font=('Arial', 32))
        self.status_indicator.pack()
        
        # 详细状态文本
        self.detailed_status_var = tk.StringVar(value="未测试")
        status_detail = ttk.Label(status_frame, textvariable=self.detailed_status_var, 
                                 wraplength=400, justify='center', font=('Arial', 10))
        status_detail.pack(pady=(0, 10))
        
        # 连接信息显示
        self.connection_info_text = scrolledtext.ScrolledText(status_frame, height=4, 
                                                             wrap=tk.WORD, font=('Consolas', 8))
        self.connection_info_text.pack(fill='both', expand=True)
        self.connection_info_text.insert(tk.END, "等待测试连接...")
        self.connection_info_text.config(state=tk.DISABLED)
        
        # === 快速配置面板 ===
        quick_config_frame = ttk.LabelFrame(parent, text="⚡ 快速配置", padding=15)
        quick_config_frame.pack(fill='x', pady=(0, 15))
        
        # 预设配置按钮 - 使用更好的样式
        presets = [
            ("OpenAI GPT-4", "openai", "🤖"),
            ("Groq Llama3", "groq", "⚡"),
            ("本地 Ollama", "ollama", "🏠")
        ]
        
        for name, preset_type, icon in presets:
            btn = ttk.Button(quick_config_frame, text=f"{icon} {name}", 
                           command=lambda p=preset_type: self.apply_preset(p))
            btn.pack(fill='x', pady=3)
        
        # === 使用统计面板 ===
        stats_frame = ttk.LabelFrame(parent, text="📊 使用统计", padding=15)
        stats_frame.pack(fill='both', expand=True, pady=(0, 15))
        
        self.stats_text = scrolledtext.ScrolledText(stats_frame, height=10, 
                                                   wrap=tk.WORD, font=('Consolas', 9))
        self.stats_text.pack(fill='both', expand=True)
        
        # 初始化统计信息 - 延迟到所有组件创建完成后
        
        # === 帮助信息面板 ===
        help_frame = ttk.LabelFrame(parent, text="💡 使用提示", padding=15)
        help_frame.pack(fill='both', expand=True)
        
        help_text = scrolledtext.ScrolledText(help_frame, height=12, 
                                            wrap=tk.WORD, font=('Arial', 9))
        help_text.pack(fill='both', expand=True)
        
        help_content = """🔧 配置步骤：
1. 选择API端点
2. 输入模型名称
3. 填写API密钥
4. 点击测试连接

🎯 推荐配置：
• OpenAI: gpt-4o, gpt-3.5-turbo
• Groq: llama3-70b-8192
• Ollama: llama3, qwen2

📄 输出格式：
• TXT: 纯文本格式，兼容性好
• DOCX: Word文档，格式丰富

⚠️ 注意事项：
• API密钥请妥善保管
• 注意请求频率限制
• 大文件建议降低并发数
• DOCX格式需要python-docx库

🔗 获取API密钥：
• OpenAI: platform.openai.com
• Groq: console.groq.com
• TogetherAI: api.together.xyz

💡 小贴士：
• 点击💡按钮查看模型建议
• 使用👁️按钮显示/隐藏密钥
• 快速配置可一键应用预设"""
        
        help_text.insert(tk.END, help_content)
        help_text.config(state=tk.DISABLED)
    
    def apply_preset(self, preset_type):
        """应用预设配置"""
        presets = {
            'openai': {
                'endpoint': 'OpenAI',
                'model': 'gpt-4o',
                'base_url': ''
            },
            'groq': {
                'endpoint': 'Groq',
                'model': 'llama3-70b-8192',
                'base_url': ''
            },
            'ollama': {
                'endpoint': 'Ollama',
                'model': 'llama3',
                'base_url': ''
            }
        }
        
        if preset_type in presets:
            preset = presets[preset_type]
            self.endpoint_var.set(preset['endpoint'])
            self.model_var.set(preset['model'])
            self.base_url_var.set(preset['base_url'])
            self.on_endpoint_change()
            self.api_status_var.set(f"✅ 已应用 {preset_type.upper()} 预设配置")
    
    def update_stats_display(self):
        """更新统计信息显示"""
        if hasattr(self, 'stats_text') and hasattr(self, 'concurrent_var'):
            stats_info = f"""📈 会话统计
━━━━━━━━━━━━━━━━━━━━
⏰ 启动时间: {time.strftime('%H:%M:%S')}
📁 扫描文件: {len(self.scanned_files)}
🔄 翻译任务: {len(self.translation_tasks)}
✅ 已完成: {sum(1 for t in self.translation_tasks.values() if t.status == '已完成')}
❌ 失败任务: {sum(1 for t in self.translation_tasks.values() if t.status == '失败')}

🔧 当前配置
━━━━━━━━━━━━━━━━━━━━
🌐 端点: {self.endpoint_var.get()}
🤖 模型: {self.model_var.get()}
🔑 密钥: {'已配置' if self.api_key_var.get() else '未配置'}
🌍 语言: {self.source_lang_var.get()} → {self.target_lang_var.get()}
⚙️ 并发数: {self.concurrent_var.get()}

💡 提示: 每30秒自动更新"""
            
            self.stats_text.config(state=tk.NORMAL)
            self.stats_text.delete(1.0, tk.END)
            self.stats_text.insert(tk.END, stats_info)
            self.stats_text.config(state=tk.DISABLED)
        
        # 每30秒更新一次统计
        self.root.after(30000, self.update_stats_display)
    
    def create_file_management_tab(self):
        """创建文件管理页面"""
        file_frame = ttk.Frame(self.notebook)
        self.notebook.add(file_frame, text="📁 文件管理")
        
        # 创建主容器，使用上下布局
        main_container = ttk.Frame(file_frame)
        main_container.pack(fill='both', expand=True, padx=10, pady=10)
        
        # 上部配置区域
        config_container = ttk.Frame(main_container)
        config_container.pack(fill='x', pady=(0, 10))
        
        # 配置区域使用两列布局
        left_config = ttk.Frame(config_container)
        left_config.pack(side='left', fill='both', expand=True, padx=(0, 10))
        
        right_config = ttk.Frame(config_container)
        right_config.pack(side='right', fill='y', padx=(10, 0))
        right_config.configure(width=300)
        right_config.pack_propagate(False)
        
        # 输入文件设置（左侧）
        input_frame = ttk.LabelFrame(left_config, text="输入文件设置", padding=10)
        input_frame.pack(fill='x', pady=(0, 5))
        
        # 输入文件夹选择
        folder_frame = ttk.Frame(input_frame)
        folder_frame.pack(fill='x', pady=2)
        
        ttk.Label(folder_frame, text="输入文件夹:").pack(side='left')
        self.input_folder_var = tk.StringVar()
        ttk.Entry(folder_frame, textvariable=self.input_folder_var, width=50).pack(
            side='left', padx=(10, 5), fill='x', expand=True)
        ttk.Button(folder_frame, text="浏览", 
                  command=self.browse_input_folder).pack(side='right')
        
        # 文件类型选择
        type_frame = ttk.Frame(input_frame)
        type_frame.pack(fill='x', pady=5)
        
        ttk.Label(type_frame, text="文件类型:").pack(side='left')
        
        self.file_types = {
            'txt': tk.BooleanVar(value=True),
            'md': tk.BooleanVar(value=True),
            'pdf': tk.BooleanVar(value=True),
            'docx': tk.BooleanVar(value=True),
            'py': tk.BooleanVar(value=False),
            'json': tk.BooleanVar(value=False),
            'cpp': tk.BooleanVar(value=False)
        }
        
        for file_type, var in self.file_types.items():
            ttk.Checkbutton(type_frame, text=file_type, variable=var).pack(side='left', padx=5)
        
        ttk.Button(input_frame, text="🔍 扫描文件", 
                  command=self.scan_files).pack(pady=5)
        
        # 输出文件设置（左侧）
        output_frame = ttk.LabelFrame(left_config, text="输出文件设置", padding=10)
        output_frame.pack(fill='x', pady=5)
        
        # 输出文件夹选择
        output_folder_frame = ttk.Frame(output_frame)
        output_folder_frame.pack(fill='x', pady=2)
        
        ttk.Label(output_folder_frame, text="输出文件夹:").pack(side='left')
        self.output_folder_var = tk.StringVar(value=str(Path.home() / "Desktop" / "translations"))
        ttk.Entry(output_folder_frame, textvariable=self.output_folder_var, width=50).pack(
            side='left', padx=(10, 5), fill='x', expand=True)
        ttk.Button(output_folder_frame, text="浏览", 
                  command=self.browse_output_folder).pack(side='right')
        
        # 输出格式选择
        format_frame = ttk.Frame(output_frame)
        format_frame.pack(fill='x', pady=5)
        
        ttk.Label(format_frame, text="输出格式:").pack(side='left')
        self.output_format_var = tk.StringVar(value="txt")
        format_combo = ttk.Combobox(format_frame, textvariable=self.output_format_var,
                                   values=["txt", "docx"], state="readonly", width=10)
        format_combo.pack(side='left', padx=(10, 20))
        
        # 格式说明
        ttk.Label(format_frame, text="• txt: 纯文本格式，兼容性好", 
                 font=('Arial', 8)).pack(side='left', padx=(10, 0))
        
        # 并发设置
        concurrent_frame = ttk.Frame(output_frame)
        concurrent_frame.pack(fill='x', pady=5)
        
        ttk.Label(concurrent_frame, text="并发任务数:", font=('Arial', 10, 'bold')).pack(side='left')
        
        # 并发数输入框
        self.concurrent_var = tk.IntVar(value=5)
        concurrent_spinbox = ttk.Spinbox(concurrent_frame, from_=1, to=10, 
                                        textvariable=self.concurrent_var, 
                                        width=5, font=('Arial', 10))
        concurrent_spinbox.pack(side='left', padx=(10, 10))
        
        # 并发数滑动条（整数步进）
        concurrent_scale = ttk.Scale(concurrent_frame, from_=1, to=10, 
                                   variable=self.concurrent_var,
                                   orient='horizontal', length=150,
                                   command=self.on_concurrent_change)
        concurrent_scale.pack(side='left', padx=(0, 10))
        
        # 并发数标签
        self.concurrent_label = ttk.Label(concurrent_frame, textvariable=self.concurrent_var,
                                         font=('Arial', 10, 'bold'), foreground='purple')
        self.concurrent_label.pack(side='left')
        
        # 并发建议
        ttk.Label(concurrent_frame, text="(建议: 小文件8-10, 大文件2-3)", 
                 font=('Arial', 8), foreground='gray').pack(side='left', padx=(10, 0))
        
        # 右侧文件预览和统计面板
        self.create_file_preview_panel(right_config)
        
        # 文件列表（下部，全宽）
        list_frame = ttk.LabelFrame(main_container, text="待翻译文件列表", padding=10)
        list_frame.pack(fill='both', expand=True)
        
        # 创建Treeview（支持多选）
        columns = ('文件名', '大小', '类型')
        self.file_tree = ttk.Treeview(list_frame, columns=columns, show='headings', height=8, selectmode='extended')
        
        for col in columns:
            self.file_tree.heading(col, text=col)
            self.file_tree.column(col, width=150)
        
        # 滚动条
        file_scrollbar = ttk.Scrollbar(list_frame, orient='vertical', command=self.file_tree.yview)
        self.file_tree.configure(yscrollcommand=file_scrollbar.set)
        
        self.file_tree.pack(side='left', fill='both', expand=True)
        file_scrollbar.pack(side='right', fill='y')
        
        # 绑定文件列表选择事件（在文件树创建后）
        self.file_tree.bind('<<TreeviewSelect>>', self.on_file_select)
        
        # 操作按钮
        action_frame = ttk.Frame(main_container)
        action_frame.pack(fill='x', pady=(10, 0))
        
        # 左侧按钮组
        left_buttons = ttk.Frame(action_frame)
        left_buttons.pack(side='left')
        
        self.start_btn = ttk.Button(left_buttons, text="🚀 开始批量翻译", 
                  command=self.start_translation)
        self.start_btn.pack(side='left', padx=(0, 10))
        
        self.pause_btn = ttk.Button(left_buttons, text="⏸️ 暂停翻译", 
                  command=self.pause_translation, state='disabled')
        self.pause_btn.pack(side='left', padx=(0, 10))
        
        self.stop_btn = ttk.Button(left_buttons, text="⏹️ 停止翻译", 
                  command=self.stop_translation, state='disabled')
        self.stop_btn.pack(side='left', padx=(0, 10))
        
        ttk.Button(left_buttons, text="🗑️ 清空任务", 
                  command=self.clear_tasks).pack(side='left', padx=(0, 10))
        
        # 右侧选择按钮组
        select_buttons = ttk.Frame(action_frame)
        select_buttons.pack(side='right')
        
        ttk.Button(select_buttons, text="✅ 全选", 
                  command=self.select_all_files, width=8).pack(side='left', padx=(0, 5))
        ttk.Button(select_buttons, text="❌ 取消全选", 
                  command=self.deselect_all_files, width=10).pack(side='left', padx=(0, 5))
        ttk.Button(select_buttons, text="🔄 反选", 
                  command=self.invert_selection, width=8).pack(side='left')
        
        # 状态显示
        self.file_status_var = tk.StringVar(value="请选择输入文件夹并扫描文件")
        ttk.Label(main_container, textvariable=self.file_status_var).pack(pady=(5, 0))
    
    def on_concurrent_change(self, value):
        """并发数改变时的处理 - 确保为整数"""
        try:
            int_value = int(float(value))
            self.concurrent_var.set(int_value)
        except (ValueError, TypeError):
            self.concurrent_var.set(5)  # 默认值
    
    def on_token_change(self, value):
        """Token数改变时的处理 - 确保为整数"""
        try:
            int_value = int(float(value))
            self.max_tokens_var.set(int_value)
        except (ValueError, TypeError):
            self.max_tokens_var.set(1000)  # 默认值
    
    def on_temp_change(self, value):
        """温度改变时的处理 - 保留一位小数"""
        try:
            float_value = round(float(value), 1)
            self.temperature_var.set(float_value)
            self.temp_label.config(text=f"{float_value:.1f}")
        except (ValueError, TypeError):
            self.temperature_var.set(0.3)  # 默认值
            self.temp_label.config(text="0.3")
    
    def on_rpm_change(self, value):
        """RPM改变时的处理 - 确保为整数"""
        try:
            int_value = int(float(value))
            self.rpm_var.set(int_value)
        except (ValueError, TypeError):
            self.rpm_var.set(60)  # 默认值
    
    def on_timeout_change(self, value):
        """超时时间改变时的处理 - 确保为整数"""
        try:
            int_value = int(float(value))
            self.api_timeout_var.set(int_value)
        except (ValueError, TypeError):
            self.api_timeout_var.set(120)  # 默认值
    
    def on_performance_mode_change(self, event=None):
        """性能模式改变时的处理"""
        mode = self.performance_mode_var.get()
        
        if mode == "快速":
            # 快速模式：适中超时，较高并发
            if hasattr(self, 'api_timeout_var'):
                self.api_timeout_var.set(180)  # 3分钟
            self.concurrent_var.set(6)
            self.rpm_var.set(80)
            if hasattr(self, 'retry_count_var'):
                self.retry_count_var.set(1)
            desc = "• 快速: 适中超时高并发，适合小文件"
        elif mode == "平衡":
            # 平衡模式：中等设置
            if hasattr(self, 'api_timeout_var'):
                self.api_timeout_var.set(300)  # 5分钟
            self.concurrent_var.set(4)
            self.rpm_var.set(60)
            if hasattr(self, 'retry_count_var'):
                self.retry_count_var.set(2)
            desc = "• 平衡: 适合大多数情况"
        elif mode == "稳定":
            # 稳定模式：较长超时，较低并发
            if hasattr(self, 'api_timeout_var'):
                self.api_timeout_var.set(600)  # 10分钟
            self.concurrent_var.set(2)
            self.rpm_var.set(30)
            if hasattr(self, 'retry_count_var'):
                self.retry_count_var.set(3)
            desc = "• 稳定: 长超时低并发，适合大文件"
        
        if hasattr(self, 'mode_desc_label'):
            self.mode_desc_label.config(text=desc)
    
    def create_file_preview_panel(self, parent):
        """创建文件预览和统计面板"""
        # 文件统计面板
        stats_frame = ttk.LabelFrame(parent, text="📊 文件统计", padding=10)
        stats_frame.pack(fill='x', pady=(0, 10))
        
        self.file_stats_text = scrolledtext.ScrolledText(stats_frame, height=6, width=30, 
                                                        wrap=tk.WORD, font=('Consolas', 9))
        self.file_stats_text.pack(fill='both', expand=True)
        
        # 文件预览面板
        preview_frame = ttk.LabelFrame(parent, text="👁️ 文件预览", padding=10)
        preview_frame.pack(fill='both', expand=True)
        
        self.file_preview_text = scrolledtext.ScrolledText(preview_frame, height=15, width=30, 
                                                          wrap=tk.WORD, font=('Arial', 9))
        self.file_preview_text.pack(fill='both', expand=True)
        
        # 绑定文件列表选择事件
        # self.file_tree.bind('<<TreeviewSelect>>', self.on_file_select)  # 移到文件树创建后
        
        # 初始化显示
        self.update_file_stats()
    
    def on_file_select(self, event):
        """文件选择事件处理"""
        if not hasattr(self, 'file_tree'):
            return
            
        selection = self.file_tree.selection()
        if selection:
            item = self.file_tree.item(selection[0])
            filename = item['values'][0]
            
            # 找到完整文件路径
            full_path = None
            for file_path in self.scanned_files:
                if os.path.basename(file_path) == filename:
                    full_path = file_path
                    break
            
            if full_path:
                self.preview_file(full_path)
    
    def preview_file(self, file_path):
        """预览文件内容"""
        if not hasattr(self, 'file_preview_text'):
            return
            
        try:
            # 读取文件前500个字符作为预览
            content = self.read_file_content(file_path)
            if content:
                preview = content[:500] + "..." if len(content) > 500 else content
                
                preview_info = f"""📄 文件: {os.path.basename(file_path)}
📁 路径: {file_path}
📏 大小: {os.path.getsize(file_path)} 字节
🔤 字符数: {len(content)}
📝 预览:
{'─' * 30}
{preview}"""
                
                self.file_preview_text.config(state=tk.NORMAL)
                self.file_preview_text.delete(1.0, tk.END)
                self.file_preview_text.insert(tk.END, preview_info)
                self.file_preview_text.config(state=tk.DISABLED)
            else:
                self.file_preview_text.config(state=tk.NORMAL)
                self.file_preview_text.delete(1.0, tk.END)
                self.file_preview_text.insert(tk.END, "❌ 无法读取文件内容")
                self.file_preview_text.config(state=tk.DISABLED)
                
        except Exception as e:
            self.file_preview_text.config(state=tk.NORMAL)
            self.file_preview_text.delete(1.0, tk.END)
            self.file_preview_text.insert(tk.END, f"❌ 预览失败: {e}")
            self.file_preview_text.config(state=tk.DISABLED)
    
    def update_file_stats(self):
        """更新文件统计信息"""
        if not hasattr(self, 'file_stats_text'):
            return
            
        total_files = len(self.scanned_files)
        total_size = sum(os.path.getsize(f) for f in self.scanned_files if os.path.exists(f))
        
        # 按类型统计
        type_stats = {}
        for file_path in self.scanned_files:
            ext = os.path.splitext(file_path)[1][1:].upper()
            type_stats[ext] = type_stats.get(ext, 0) + 1
        
        avg_size = (total_size/total_files/1024) if total_files > 0 else 0
        
        stats_info = f"""📈 文件统计
━━━━━━━━━━━━━━━━━━━━
📁 总文件数: {total_files}
💾 总大小: {total_size/1024:.1f} KB
📊 平均大小: {avg_size:.1f} KB

📋 类型分布:
"""
        
        for file_type, count in sorted(type_stats.items()):
            stats_info += f"• {file_type}: {count} 个\n"
        
        if not self.scanned_files:
            stats_info = "📂 暂无文件\n请先扫描文件夹"
        
        self.file_stats_text.config(state=tk.NORMAL)
        self.file_stats_text.delete(1.0, tk.END)
        self.file_stats_text.insert(tk.END, stats_info)
        self.file_stats_text.config(state=tk.DISABLED)
    
    def create_progress_tab(self):
        """创建翻译进度页面"""
        progress_frame = ttk.Frame(self.notebook)
        self.notebook.add(progress_frame, text="📊 翻译进度")
        
        # 创建主容器
        main_container = ttk.Frame(progress_frame)
        main_container.pack(fill='both', expand=True, padx=10, pady=10)
        
        # 上部总体进度区域
        summary_container = ttk.Frame(main_container)
        summary_container.pack(fill='x', pady=(0, 10))
        
        # 总体进度（左侧）
        summary_frame = ttk.LabelFrame(summary_container, text="总体进度", padding=10)
        summary_frame.pack(side='left', fill='both', expand=True, padx=(0, 10))
        
        self.overall_progress_var = tk.StringVar(value="暂无翻译任务")
        ttk.Label(summary_frame, textvariable=self.overall_progress_var, 
                 style='Heading.TLabel').pack()
        
        self.overall_progress_bar = Progressbar(summary_frame, length=400, mode='determinate')
        self.overall_progress_bar.pack(pady=5)
        
        # 实时统计（右侧）
        realtime_frame = ttk.LabelFrame(summary_container, text="📈 实时统计", padding=10)
        realtime_frame.pack(side='right', fill='y', padx=(10, 0))
        realtime_frame.configure(width=250)
        realtime_frame.pack_propagate(False)
        
        self.realtime_stats_text = scrolledtext.ScrolledText(realtime_frame, height=6, width=25, 
                                                            wrap=tk.WORD, font=('Consolas', 9))
        self.realtime_stats_text.pack(fill='both', expand=True)
        
        # 下部详细进度区域
        detail_container = ttk.Frame(main_container)
        detail_container.pack(fill='both', expand=True)
        
        # 详细进度列表（左侧）
        detail_frame = ttk.LabelFrame(detail_container, text="详细进度", padding=10)
        detail_frame.pack(side='left', fill='both', expand=True, padx=(0, 10))
        
        # 创建进度Treeview
        progress_columns = ('文件名', '状态', '进度', '耗时', '错误信息')
        self.progress_tree = ttk.Treeview(detail_frame, columns=progress_columns, show='headings')
        
        for col in progress_columns:
            self.progress_tree.heading(col, text=col)
            if col == '错误信息':
                self.progress_tree.column(col, width=200)
            else:
                self.progress_tree.column(col, width=120)
        
        # 滚动条
        progress_scrollbar = ttk.Scrollbar(detail_frame, orient='vertical', command=self.progress_tree.yview)
        self.progress_tree.configure(yscrollcommand=progress_scrollbar.set)
        
        self.progress_tree.pack(side='left', fill='both', expand=True)
        progress_scrollbar.pack(side='right', fill='y')
        
        # 任务详情面板（右侧）- 增加宽度以更好显示翻译内容
        task_detail_frame = ttk.LabelFrame(detail_container, text="🔍 任务详情", padding=10)
        task_detail_frame.pack(side='right', fill='both', expand=True, padx=(10, 0))
        
        # 添加操作按钮
        detail_button_frame = ttk.Frame(task_detail_frame)
        detail_button_frame.pack(fill='x', pady=(0, 5))
        
        ttk.Button(detail_button_frame, text="📋 复制最终翻译", 
                  command=self.copy_final_translation, width=15).pack(side='left', padx=(0, 5))
        ttk.Button(detail_button_frame, text="📋 复制初始翻译", 
                  command=self.copy_init_translation, width=15).pack(side='left', padx=(0, 5))
        ttk.Button(detail_button_frame, text="📋 复制反思建议", 
                  command=self.copy_reflect_translation, width=15).pack(side='left')
        
        self.task_detail_text = scrolledtext.ScrolledText(task_detail_frame, height=20, 
                                                         wrap=tk.WORD, font=('Consolas', 9))
        self.task_detail_text.pack(fill='both', expand=True)
        
        # 绑定进度列表选择事件
        self.progress_tree.bind('<<TreeviewSelect>>', self.on_task_select)
        
        # 操作按钮
        progress_button_frame = ttk.Frame(main_container)
        progress_button_frame.pack(fill='x', pady=(10, 0))
        
        ttk.Button(progress_button_frame, text="🔄 刷新进度", 
                  command=self.refresh_progress).pack(side='left', padx=(0, 10))
        ttk.Button(progress_button_frame, text="📁 打开输出文件夹", 
                  command=self.open_output_folder).pack(side='left')
        
        # 启动进度更新定时器
        self.update_progress_display()
        
        # 启动统计更新定时器（在所有组件创建完成后）
        self.root.after(1000, self.update_stats_display)  # 延迟1秒启动
    
    def on_task_select(self, event):
        """任务选择事件处理"""
        selection = self.progress_tree.selection()
        if selection:
            item = self.progress_tree.item(selection[0])
            filename = item['values'][0]
            
            # 找到对应的任务
            task = None
            for t in self.translation_tasks.values():
                if t.filename == filename:
                    task = t
                    break
            
            if task:
                self.show_task_detail(task)
    
    def show_task_detail(self, task):
        """显示任务详情 - 使用标签页展示翻译结果"""
        # 保存当前选中的任务（用于复制功能）
        self.current_selected_task = task
        
        # 清空当前内容
        self.task_detail_text.config(state=tk.NORMAL)
        self.task_detail_text.delete(1.0, tk.END)
        
        # 基本信息
        basic_info = f"""📋 任务详情
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📄 文件名: {task.filename}
🆔 任务ID: {task.task_id}
📊 状态: {task.status}
📈 进度: {task.progress}%

⏱️ 时间信息:
• 开始: {time.strftime('%H:%M:%S', time.localtime(task.start_time)) if task.start_time else '未开始'}
• 结束: {time.strftime('%H:%M:%S', time.localtime(task.end_time)) if task.end_time else '进行中'}
• 耗时: {f'{(task.end_time - task.start_time):.1f}秒' if task.start_time and task.end_time else '计算中'}

📝 内容统计:
• 原文: {len(task.content)} 字符
• 初始翻译: {len(task.init_translation)} 字符
• 最终翻译: {len(task.final_translation)} 字符

"""
        
        # 根据任务状态显示不同内容
        if task.status == "已完成":
            # 显示翻译结果
            detail_content = basic_info + f"""━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📌 最终翻译结果:
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
{task.final_translation}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📝 初始翻译:
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
{task.init_translation}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
💭 反思建议:
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
{task.reflect_translation}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
💡 提示: 
• 双击文本可选择复制
• 翻译结果已自动保存到输出文件夹
• 点击"📁 打开输出文件夹"查看文件"""
        
        elif task.status == "失败":
            # 显示错误信息
            detail_content = basic_info + f"""━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
❌ 错误信息:
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
{task.error_message}

💡 可能的解决方案:
• 检查API配置是否正确
• 确认网络连接正常
• 查看max_tokens设置是否足够
• 检查文件内容是否有特殊字符
• 尝试降低并发任务数"""
        
        elif task.status == "翻译中":
            # 显示进度信息
            elapsed = time.time() - task.start_time if task.start_time else 0
            detail_content = basic_info + f"""━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🔄 翻译进行中...
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
• 已用时间: {elapsed:.1f}秒
• 当前进度: {task.progress}%
• 预计剩余: 计算中...

💡 翻译流程:
1. ✅ 模型加载 (10%)
2. {'✅' if task.progress > 20 else '⏳'} 初始翻译 (20-60%)
3. {'✅' if task.progress > 60 else '⏳'} 反思评估 (60-80%)
4. {'✅' if task.progress > 80 else '⏳'} 改进翻译 (80-100%)

请耐心等待..."""
        
        else:
            # 等待中
            detail_content = basic_info + f"""━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
⏳ 等待翻译...
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
任务已加入队列，等待处理

💡 提示:
• 当前可能有其他任务正在翻译
• 并发任务数: {self.concurrent_var.get() if hasattr(self, 'concurrent_var') else '未知'}
• 请耐心等待"""
        
        self.task_detail_text.insert(tk.END, detail_content)
        self.task_detail_text.config(state=tk.DISABLED)
    
    def copy_final_translation(self):
        """复制最终翻译到剪贴板"""
        if hasattr(self, 'current_selected_task') and self.current_selected_task:
            task = self.current_selected_task
            if task.final_translation:
                self.root.clipboard_clear()
                self.root.clipboard_append(task.final_translation)
                messagebox.showinfo("成功", "最终翻译已复制到剪贴板")
            else:
                messagebox.showwarning("提示", "该任务还没有最终翻译结果")
        else:
            messagebox.showwarning("提示", "请先选择一个任务")
    
    def copy_init_translation(self):
        """复制初始翻译到剪贴板"""
        if hasattr(self, 'current_selected_task') and self.current_selected_task:
            task = self.current_selected_task
            if task.init_translation:
                self.root.clipboard_clear()
                self.root.clipboard_append(task.init_translation)
                messagebox.showinfo("成功", "初始翻译已复制到剪贴板")
            else:
                messagebox.showwarning("提示", "该任务还没有初始翻译结果")
        else:
            messagebox.showwarning("提示", "请先选择一个任务")
    
    def copy_reflect_translation(self):
        """复制反思建议到剪贴板"""
        if hasattr(self, 'current_selected_task') and self.current_selected_task:
            task = self.current_selected_task
            if task.reflect_translation:
                self.root.clipboard_clear()
                self.root.clipboard_append(task.reflect_translation)
                messagebox.showinfo("成功", "反思建议已复制到剪贴板")
            else:
                messagebox.showwarning("提示", "该任务还没有反思建议")
        else:
            messagebox.showwarning("提示", "请先选择一个任务")
    
    def update_realtime_stats(self):
        """更新实时统计"""
        if hasattr(self, 'realtime_stats_text'):
            total_tasks = len(self.translation_tasks)
            completed = sum(1 for t in self.translation_tasks.values() if t.status == '已完成')
            failed = sum(1 for t in self.translation_tasks.values() if t.status == '失败')
            in_progress = sum(1 for t in self.translation_tasks.values() if t.status == '翻译中')
            waiting = total_tasks - completed - failed - in_progress
            
            # 计算平均耗时
            completed_tasks = [t for t in self.translation_tasks.values() 
                             if t.status == '已完成' and t.start_time and t.end_time]
            avg_time = sum(t.end_time - t.start_time for t in completed_tasks) / len(completed_tasks) if completed_tasks else 0
            
            stats_info = f"""⏰ {time.strftime('%H:%M:%S')}
━━━━━━━━━━━━━━━━━━━━
📊 任务统计:
• 总任务: {total_tasks}
• ✅ 已完成: {completed}
• 🔄 进行中: {in_progress}
• ⏳ 等待中: {waiting}
• ❌ 失败: {failed}

📈 性能指标:
• 完成率: {f'{(completed/total_tasks*100):.1f}%' if total_tasks > 0 else '0%'}
• 平均耗时: {avg_time:.1f}s
• 预计剩余: {f'{(avg_time * (total_tasks - completed)):.0f}s' if avg_time > 0 and total_tasks > completed else '0s'}

🔄 状态: {'翻译中' if self.is_translating else '空闲'}"""
            
            self.realtime_stats_text.config(state=tk.NORMAL)
            self.realtime_stats_text.delete(1.0, tk.END)
            self.realtime_stats_text.insert(tk.END, stats_info)
            self.realtime_stats_text.config(state=tk.DISABLED)
    
    def create_about_tab(self):
        """创建关于页面"""
        about_frame = ttk.Frame(self.notebook)
        self.notebook.add(about_frame, text="ℹ️ 关于")
        
        # 创建滚动文本
        about_text = scrolledtext.ScrolledText(about_frame, wrap=tk.WORD, width=80, height=30)
        about_text.pack(fill='both', expand=True, padx=10, pady=10)
        
        about_content = """
🚀 Translation Agent Pro - 专业批量翻译软件

版本: 2.0.0
作者: Translation Agent Team
许可: MIT License

✨ 主要功能

• 🔧 API配置管理: 支持多种LLM提供商，可测试连接状态
• 📁 智能文件管理: 批量扫描文件夹，支持多种文件格式
• 🚀 并发翻译处理: 最大10个文件同时翻译，提高效率
• 📊 实时进度监控: 详细的翻译进度显示和状态跟踪
• 💾 自动结果保存: 翻译完成后自动保存到指定文件夹
• 🎯 反思式翻译: 三阶段翻译流程，确保翻译质量

🔧 支持的API端点

• OpenAI: GPT-4, GPT-3.5等模型
• Groq: Llama3, Mixtral等高速推理
• TogetherAI: 开源模型集合
• Ollama: 本地部署模型
• CUSTOM: 自定义OpenAI兼容端点

📄 支持的文件格式

• 文本文件: .txt, .md
• 文档文件: .pdf, .docx
• 代码文件: .py, .json, .cpp

🎯 使用流程

1. 配置API: 在"API设置"页面配置你的API密钥和参数
2. 测试连接: 点击"测试API连接"确保配置正确
3. 选择文件: 在"文件管理"页面选择输入和输出文件夹
4. 开始翻译: 点击"开始批量翻译"启动任务
5. 监控进度: 在"翻译进度"页面查看实时进度

📞 技术支持

如遇问题，请检查：
• API密钥是否正确
• 网络连接是否正常
• 文件格式是否支持
• 输出文件夹是否有写入权限

🎯 翻译质量保证

本软件采用反思式翻译工作流：
1. 初始翻译: 快速生成初步翻译
2. 反思评估: 从准确性、流畅性、风格、术语四个维度评估
3. 改进翻译: 根据反思建议优化翻译

这种方法类似于人类翻译的"初译→审校→定稿"流程，显著提升翻译质量。

© 2024 Translation Agent Team. All rights reserved.
        """
        
        about_text.insert(tk.END, about_content)
        about_text.config(state=tk.DISABLED)
    
    def show_model_suggestions(self):
        """显示模型建议"""
        endpoint = self.endpoint_var.get()
        suggestions = {
            "OpenAI": ["gpt-4o", "gpt-4o-mini", "gpt-3.5-turbo", "gpt-4-turbo"],
            "Groq": ["llama3-70b-8192", "llama3-8b-8192", "mixtral-8x7b-32768", "gemma-7b-it"],
            "TogetherAI": ["Qwen/Qwen2-72B-Instruct", "meta-llama/Llama-2-70b-chat-hf", "mistralai/Mixtral-8x7B-Instruct-v0.1"],
            "Ollama": ["llama3", "qwen2", "mistral", "codellama"],
            "CUSTOM": ["请输入自定义模型名称"]
        }
        
        models = suggestions.get(endpoint, [])
        suggestion_text = f"{endpoint} 推荐模型:\n" + "\n".join(f"• {model}" for model in models)
        
        messagebox.showinfo("模型建议", suggestion_text)
    
    def on_endpoint_change(self, event=None):
        """端点改变时的处理"""
        endpoint = self.endpoint_var.get()
        
        # 更新端点状态指示器
        status_colors = {
            "OpenAI": "🟢",
            "Groq": "🟡", 
            "TogetherAI": "🔵",
            "Ollama": "🟠",
            "CUSTOM": "🟣"
        }
        self.endpoint_status.config(text=status_colors.get(endpoint, "⚪"))
        
        # 如果正在加载配置，不更新模型名（保持加载的值）
        if not self.is_loading_config:
            # 更新默认模型
            model_map = {
                "OpenAI": "gpt-4o",
                "Groq": "llama3-70b-8192",
                "TogetherAI": "Qwen/Qwen2-72B-Instruct",
                "Ollama": "llama3",
                "CUSTOM": ""
            }
            
            if endpoint in model_map:
                self.model_var.set(model_map[endpoint])
        
        # 显示/隐藏基础URL
        if endpoint == "CUSTOM":
            self.base_url_row.pack(fill='x', pady=(0, 12))
        else:
            self.base_url_row.pack_forget()
    
    def on_endpoint2_change(self, event=None):
        """额外端点改变时的处理"""
        endpoint2 = self.endpoint2_var.get()
        
        # 如果正在加载配置，不更新模型名（保持加载的值）
        if not self.is_loading_config:
            # 更新默认模型
            model_map = {
                "OpenAI": "gpt-4o",
                "Groq": "llama3-70b-8192",
                "TogetherAI": "Qwen/Qwen2-72B-Instruct",
                "Ollama": "llama3",
                "CUSTOM": ""
            }
            
            if endpoint2 in model_map:
                self.model2_var.set(model_map[endpoint2])
        
        # 显示/隐藏额外基础URL
        if endpoint2 == "CUSTOM":
            self.base_url2_row.pack(fill='x', pady=(0, 10))
        else:
            self.base_url2_row.pack_forget()
    
    def toggle_extra_endpoint(self):
        """切换额外端点显示"""
        if self.use_extra_endpoint_var.get():
            self.extra_endpoint_frame.pack(fill='x', pady=(10, 0))
        else:
            self.extra_endpoint_frame.pack_forget()
    
    def browse_input_folder(self):
        """浏览输入文件夹"""
        folder = filedialog.askdirectory(title="选择输入文件夹")
        if folder:
            self.input_folder_var.set(folder)
    
    def browse_output_folder(self):
        """浏览输出文件夹"""
        folder = filedialog.askdirectory(title="选择输出文件夹")
        if folder:
            self.output_folder_var.set(folder)
    
    def scan_files(self):
        """扫描文件"""
        input_folder = self.input_folder_var.get()
        if not input_folder or not os.path.exists(input_folder):
            messagebox.showerror("错误", "请选择有效的输入文件夹")
            return
        
        # 获取选中的文件类型
        selected_types = [ext for ext, var in self.file_types.items() if var.get()]
        if not selected_types:
            messagebox.showerror("错误", "请至少选择一种文件类型")
            return
        
        # 扫描文件
        self.scanned_files = []
        for ext in selected_types:
            pattern = os.path.join(input_folder, f"*.{ext}")
            import glob
            self.scanned_files.extend(glob.glob(pattern))
        
        # 按文件名自然排序（支持章节号）
        self.scanned_files = self.natural_sort_files(self.scanned_files)
        
        # 更新文件列表（如果file_tree已创建）
        if hasattr(self, 'file_tree'):
            self.update_file_list()
        
        # 更新统计信息
        self.update_file_stats()
        
        if self.scanned_files:
            self.file_status_var.set(f"找到 {len(self.scanned_files)} 个文件")
        else:
            self.file_status_var.set("未找到匹配的文件")
    
    def natural_sort_files(self, file_list):
        """自然排序文件列表（支持数字章节排序）"""
        import re
        
        def natural_sort_key(file_path):
            """生成自然排序的键"""
            filename = os.path.basename(file_path)
            # 将文件名中的数字转换为整数进行排序
            parts = re.split(r'(\d+)', filename)
            return [int(part) if part.isdigit() else part.lower() for part in parts]
        
        return sorted(file_list, key=natural_sort_key)
    
    def select_all_files(self):
        """全选文件"""
        if not hasattr(self, 'file_tree'):
            return
        
        # 选择所有项目
        all_items = self.file_tree.get_children()
        self.file_tree.selection_set(all_items)
        
        # 更新状态
        if all_items:
            self.file_status_var.set(f"已选择 {len(all_items)} 个文件")
    
    def deselect_all_files(self):
        """取消全选"""
        if not hasattr(self, 'file_tree'):
            return
        
        # 清除所有选择
        self.file_tree.selection_remove(self.file_tree.selection())
        
        # 更新状态
        self.file_status_var.set(f"已取消选择，共 {len(self.scanned_files)} 个文件")
    
    def invert_selection(self):
        """反选"""
        if not hasattr(self, 'file_tree'):
            return
        
        # 获取所有项目和当前选择
        all_items = self.file_tree.get_children()
        selected_items = set(self.file_tree.selection())
        
        # 计算反选项目
        inverted_items = [item for item in all_items if item not in selected_items]
        
        # 清除当前选择并设置新选择
        self.file_tree.selection_remove(self.file_tree.selection())
        self.file_tree.selection_set(inverted_items)
        
        # 更新状态
        if inverted_items:
            self.file_status_var.set(f"已选择 {len(inverted_items)} 个文件")
    
    def update_file_list(self):
        """更新文件列表显示"""
        # 检查file_tree是否存在
        if not hasattr(self, 'file_tree'):
            return
            
        # 清空现有项目
        for item in self.file_tree.get_children():
            self.file_tree.delete(item)
        
        # 添加文件
        for file_path in self.scanned_files:
            filename = os.path.basename(file_path)
            file_size = os.path.getsize(file_path)
            size_str = f"{file_size/1024:.1f} KB" if file_size < 1024*1024 else f"{file_size/(1024*1024):.1f} MB"
            file_type = os.path.splitext(filename)[1][1:].upper()
            
            self.file_tree.insert('', 'end', values=(filename, size_str, file_type))
    
    def test_api_connection(self):
        """测试API连接"""
        try:
            endpoint = self.endpoint_var.get()
            model = self.model_var.get()
            api_key = self.api_key_var.get()
            base_url = self.base_url_var.get()
            
            if not api_key:
                self.api_status_var.set("❌ 请输入API密钥")
                return
            
            if not model:
                self.api_status_var.set("❌ 请输入模型名称")
                return
            
            # 显示测试中状态
            self.api_status_var.set("🔍 正在测试API连接...")
            if hasattr(self, 'status_indicator'):
                self.status_indicator.config(text="🟡")
                self.detailed_status_var.set("正在测试连接...")
                
                # 更新连接信息
                if hasattr(self, 'connection_info_text'):
                    self.connection_info_text.config(state=tk.NORMAL)
                    self.connection_info_text.delete(1.0, tk.END)
                    self.connection_info_text.insert(tk.END, f"测试端点: {endpoint}\n模型: {model}\n状态: 连接中...")
                    self.connection_info_text.config(state=tk.DISABLED)
                    
            self.root.update()
            
            # 在后台线程中测试
            def test_thread():
                try:
                    print(f"\n[API测试] 开始测试连接")
                    print(f"[API测试] 端点: {endpoint}")
                    print(f"[API测试] 模型: {model}")
                    print(f"[API测试] 基础URL: {base_url}")
                    
                    # 使用更简单的测试方法
                    import concurrent.futures
                    import requests
                    import json
                    
                    def simple_api_test():
                        """简单的API测试"""
                        # 构建API请求
                        if endpoint == "CUSTOM":
                            if not base_url:
                                raise Exception("CUSTOM端点需要设置基础URL")
                            
                            # 确保URL格式正确
                            test_url = base_url.rstrip('/')
                            if not test_url.endswith('/v1'):
                                test_url += '/v1'
                            test_url += '/chat/completions'
                        else:
                            # 其他端点的默认URL
                            endpoint_urls = {
                                "OpenAI": "https://api.openai.com/v1/chat/completions",
                                "Groq": "https://api.groq.com/openai/v1/chat/completions",
                                "TogetherAI": "https://api.together.xyz/v1/chat/completions"
                            }
                            test_url = endpoint_urls.get(endpoint)
                            if not test_url:
                                raise Exception(f"不支持的端点: {endpoint}")
                        
                        # 构建请求头
                        headers = {
                            "Content-Type": "application/json",
                            "Authorization": f"Bearer {api_key}"
                        }
                        
                        # 构建请求数据
                        data = {
                            "model": model,
                            "messages": [
                                {"role": "user", "content": "Hello"}
                            ],
                            "max_tokens": 50,
                            "temperature": 0.1
                        }
                        
                        print(f"[API测试] 请求URL: {test_url}")
                        print(f"[API测试] 发送请求...")
                        
                        # 发送请求（15秒超时）
                        response = requests.post(
                            test_url,
                            headers=headers,
                            json=data,
                            timeout=15
                        )
                        
                        print(f"[API测试] 响应状态码: {response.status_code}")
                        
                        if response.status_code == 200:
                            result = response.json()
                            if 'choices' in result and len(result['choices']) > 0:
                                content = result['choices'][0]['message']['content']
                                print(f"[API测试] 响应内容: {content}")
                                return True, "连接成功"
                            else:
                                return False, "API响应格式异常"
                        else:
                            error_text = response.text
                            print(f"[API测试] 错误响应: {error_text}")
                            return False, f"HTTP {response.status_code}: {error_text[:200]}"
                    
                    # 使用线程池执行器设置20秒超时
                    with concurrent.futures.ThreadPoolExecutor(max_workers=1) as executor:
                        future = executor.submit(simple_api_test)
                        try:
                            success, message = future.result(timeout=20.0)
                        except concurrent.futures.TimeoutError:
                            raise Exception("API响应超时（20秒）：请检查网络连接和API服务器状态")
                    
                    if success:
                        print(f"[API测试] ✅ 测试成功")
                        self.root.after(0, lambda: self.api_status_var.set("✅ API连接测试成功！"))
                        if hasattr(self, 'status_indicator'):
                            self.root.after(0, lambda: self.status_indicator.config(text="🟢"))
                            self.root.after(0, lambda: self.detailed_status_var.set("连接正常\n可以开始翻译"))
                            
                            # 更新连接信息
                            if hasattr(self, 'connection_info_text'):
                                def update_success_info():
                                    self.connection_info_text.config(state=tk.NORMAL)
                                    self.connection_info_text.delete(1.0, tk.END)
                                    self.connection_info_text.insert(tk.END, f"✅ 连接成功\n端点: {endpoint}\n模型: {model}\n响应: 正常\n测试: 通过")
                                    self.connection_info_text.config(state=tk.DISABLED)
                                self.root.after(0, update_success_info)
                    else:
                        print(f"[API测试] ❌ 测试失败: {message}")
                        self.root.after(0, lambda: self.api_status_var.set(f"❌ {message}"))
                        if hasattr(self, 'status_indicator'):
                            self.root.after(0, lambda: self.status_indicator.config(text="🔴"))
                            self.root.after(0, lambda: self.detailed_status_var.set(f"连接失败\n{message[:50]}"))
                            
                            # 更新连接信息
                            if hasattr(self, 'connection_info_text'):
                                def update_error_info():
                                    self.connection_info_text.config(state=tk.NORMAL)
                                    self.connection_info_text.delete(1.0, tk.END)
                                    self.connection_info_text.insert(tk.END, f"❌ 连接失败\n端点: {endpoint}\n模型: {model}\n错误: {message}")
                                    self.connection_info_text.config(state=tk.DISABLED)
                                self.root.after(0, update_error_info)
                        
                except Exception as e:
                    error_msg = str(e)
                    print(f"[API测试] ❌ 异常: {error_msg}")
                    
                    if "timeout" in error_msg.lower():
                        msg = "❌ 连接超时: 请检查网络连接"
                        detail = "连接超时\n检查网络"
                    elif "401" in error_msg:
                        msg = "❌ API密钥无效 (401): 请检查API密钥"
                        detail = "401错误\n密钥无效"
                    elif "404" in error_msg:
                        msg = "❌ API端点错误 (404): 请检查基础URL和模型名称"
                        detail = "404错误\n检查URL和模型"
                    elif "connection" in error_msg.lower():
                        msg = "❌ 网络连接失败: 请检查网络设置"
                        detail = "网络连接失败"
                    else:
                        msg = f"❌ 连接失败: {error_msg[:50]}"
                        detail = f"连接失败\n{error_msg[:30]}"
                    
                    self.root.after(0, lambda: self.api_status_var.set(msg))
                    if hasattr(self, 'status_indicator'):
                        self.root.after(0, lambda: self.status_indicator.config(text="🔴"))
                        self.root.after(0, lambda: self.detailed_status_var.set(detail))
                        
                        # 更新连接信息
                        if hasattr(self, 'connection_info_text'):
                            def update_error_info():
                                self.connection_info_text.config(state=tk.NORMAL)
                                self.connection_info_text.delete(1.0, tk.END)
                                self.connection_info_text.insert(tk.END, f"❌ 连接失败\n端点: {endpoint}\n模型: {model}\n错误: {error_msg}")
                                self.connection_info_text.config(state=tk.DISABLED)
                            self.root.after(0, update_error_info)
            
            threading.Thread(target=test_thread, daemon=True).start()
            
        except Exception as e:
            print(f"[API测试] 启动测试失败: {e}")
            self.api_status_var.set(f"❌ 测试失败: {e}")
    
    def start_translation(self):
        """开始翻译"""
        if self.is_translating:
            messagebox.showwarning("警告", "翻译正在进行中，请等待完成")
            return
        
        if not self.scanned_files:
            messagebox.showerror("错误", "请先扫描文件")
            return
        
        if not self.api_key_var.get():
            messagebox.showerror("错误", "请先配置API密钥")
            return
        
        output_folder = self.output_folder_var.get()
        if not output_folder:
            messagebox.showerror("错误", "请选择输出文件夹")
            return
        
        # 创建输出文件夹
        try:
            os.makedirs(output_folder, exist_ok=True)
        except Exception as e:
            messagebox.showerror("错误", f"无法创建输出文件夹: {e}")
            return
        
        # 获取选中的文件
        selected_items = self.file_tree.selection() if hasattr(self, 'file_tree') else []
        
        if not selected_items:
            # 如果没有选中任何文件，询问是否翻译所有文件
            if not messagebox.askyesno("确认", "未选择任何文件，是否翻译所有文件？"):
                return
            selected_files = self.scanned_files
        else:
            # 获取选中文件的路径
            selected_files = []
            for item in selected_items:
                filename = self.file_tree.item(item)['values'][0]
                # 从scanned_files中找到对应的完整路径
                for file_path in self.scanned_files:
                    if os.path.basename(file_path) == filename:
                        selected_files.append(file_path)
                        break
        
        if not selected_files:
            messagebox.showerror("错误", "没有可翻译的文件")
            return
        
        # 准备翻译任务
        self.translation_tasks.clear()
        self.task_counter = 0
        
        for file_path in selected_files:
            try:
                # 读取文件内容
                content = self.read_file_content(file_path)
                if content:
                    self.task_counter += 1
                    task_id = f"task_{self.task_counter}"
                    filename = os.path.splitext(os.path.basename(file_path))[0]
                    task = TranslationTask(task_id, filename, content, file_path)
                    self.translation_tasks[task_id] = task
            except Exception as e:
                print(f"读取文件 {file_path} 失败: {e}")
        
        if not self.translation_tasks:
            messagebox.showerror("错误", "没有可翻译的文件内容")
            return
        
        # 开始翻译
        self.is_translating = True
        self.is_paused = False
        self.file_status_var.set(f"开始翻译 {len(self.translation_tasks)} 个文件...")
        
        # 更新按钮状态
        self.start_btn.config(state='disabled')
        self.pause_btn.config(state='normal', text="⏸️ 暂停翻译")
        self.stop_btn.config(state='normal')
        
        # 在后台线程中执行翻译
        threading.Thread(target=self.run_translation, daemon=True).start()
    
    def pause_translation(self):
        """暂停/继续翻译"""
        if self.is_paused:
            # 继续翻译
            self.is_paused = False
            self.pause_btn.config(text="⏸️ 暂停翻译")
            self.file_status_var.set("继续翻译...")
        else:
            # 暂停翻译
            self.is_paused = True
            self.pause_btn.config(text="▶️ 继续翻译")
            self.file_status_var.set("翻译已暂停")
    
    def stop_translation(self):
        """停止翻译"""
        if messagebox.askyesno("确认", "确定要停止翻译吗？未完成的任务将被取消。"):
            self.is_translating = False
            self.is_paused = False
            self.file_status_var.set("翻译已停止")
            
            # 更新按钮状态
            self.start_btn.config(state='normal')
            self.pause_btn.config(state='disabled', text="⏸️ 暂停翻译")
            self.stop_btn.config(state='disabled')
            
            # 标记未完成的任务为已取消
            for task in self.translation_tasks.values():
                if task.status == "翻译中" or task.status == "等待中":
                    task.status = "已取消"
                    task.error_message = "用户停止翻译"
    
    def read_file_content(self, file_path):
        """读取文件内容"""
        try:
            file_ext = os.path.splitext(file_path)[1].lower()
            
            if file_ext == '.pdf':
                return extract_pdf(file_path)
            elif file_ext == '.docx':
                return extract_docx(file_path)
            else:
                return extract_text(file_path)
        except Exception as e:
            print(f"读取文件内容失败: {e}")
            return None
    
    def run_translation(self):
        """运行翻译任务 - 优化版本，增加更好的并发控制和错误处理"""
        try:
            print(f"\n{'#'*60}")
            print(f"# 开始批量翻译")
            print(f"# 总任务数: {len(self.translation_tasks)}")
            print(f"# 并发数: {self.concurrent_var.get()}")
            print(f"{'#'*60}\n")
            
            # 获取配置
            config = {
                'endpoint': self.endpoint_var.get(),
                'model': self.model_var.get(),
                'api_key': self.api_key_var.get(),
                'base_url': self.base_url_var.get(),
                'use_extra_endpoint': self.use_extra_endpoint_var.get(),
                'endpoint2': self.endpoint2_var.get(),
                'model2': self.model2_var.get(),
                'api_key2': self.api_key2_var.get(),
                'base_url2': self.base_url2_var.get(),
                'source_lang': self.source_lang_var.get(),
                'target_lang': self.target_lang_var.get(),
                'country': self.country_var.get(),
                'max_tokens': self.max_tokens_var.get(),
                'temperature': self.temperature_var.get(),
                'rpm': self.rpm_var.get()
            }
            
            concurrent_tasks = self.concurrent_var.get()
            output_folder = self.output_folder_var.get()
            
            completed_count = 0
            failed_count = 0
            
            # 将任务列表转换为队列，便于管理
            task_queue = list(self.translation_tasks.values())
            active_futures = {}
            
            # 使用线程池执行翻译 - 优化版本
            with ThreadPoolExecutor(max_workers=concurrent_tasks) as executor:
                # 提交初始批次的任务
                for i in range(min(concurrent_tasks, len(task_queue))):
                    if not self.is_translating:
                        break
                    
                    task = task_queue.pop(0)
                    print(f"[任务管理] 提交任务: {task.filename}")
                    future = executor.submit(self.translate_single_file, task, config)
                    active_futures[future] = task
                
                # 处理完成的任务并提交新任务
                while active_futures and self.is_translating:
                    # 等待暂停结束
                    while self.is_paused and self.is_translating:
                        time.sleep(0.5)
                    
                    if not self.is_translating:
                        print("⚠️ 用户停止翻译，取消剩余任务")
                        break
                    
                    # 检查已完成的任务（设置较短的超时以便及时响应暂停/停止）
                    try:
                        completed_futures = []
                        for future in list(active_futures.keys()):
                            if future.done():
                                completed_futures.append(future)
                        
                        # 如果没有完成的任务，短暂等待
                        if not completed_futures:
                            time.sleep(0.1)
                            continue
                        
                        # 处理完成的任务
                        for future in completed_futures:
                            task = active_futures.pop(future)
                            
                            try:
                                completed_task = future.result()
                                if completed_task.status == "已完成":
                                    completed_count += 1
                                    print(f"✓ 任务完成: {completed_task.filename}")
                                    self.save_translation_result(completed_task, output_folder)
                                elif completed_task.status == "失败":
                                    failed_count += 1
                                    print(f"✗ 任务失败: {completed_task.filename} - {completed_task.error_message}")
                                else:
                                    print(f"⚠️ 任务状态异常: {completed_task.filename} - {completed_task.status}")
                                    
                            except Exception as e:
                                failed_count += 1
                                task.status = "失败"
                                task.error_message = f"处理任务结果时出错: {str(e)}"
                                print(f"✗ 处理任务结果时出错: {task.filename} - {str(e)}")
                                import traceback
                                traceback.print_exc()
                        
                        # 提交新任务（如果还有待处理的任务）
                        while len(active_futures) < concurrent_tasks and task_queue and self.is_translating:
                            if self.is_paused:
                                break
                                
                            task = task_queue.pop(0)
                            print(f"[任务管理] 提交新任务: {task.filename}")
                            future = executor.submit(self.translate_single_file, task, config)
                            active_futures[future] = task
                        
                        # 更新状态显示
                        total_tasks = len(self.translation_tasks)
                        processed = completed_count + failed_count
                        remaining = total_tasks - processed
                        
                        status_msg = f"翻译进度: {processed}/{total_tasks} (成功:{completed_count}, 失败:{failed_count}, 剩余:{remaining})"
                        self.root.after(0, lambda msg=status_msg: self.file_status_var.set(msg))
                        
                    except Exception as e:
                        print(f"⚠️ 任务管理循环中出错: {e}")
                        import traceback
                        traceback.print_exc()
                        time.sleep(1)  # 出错时稍作等待
                
                # 等待所有剩余任务完成
                print(f"[任务管理] 等待剩余 {len(active_futures)} 个任务完成...")
                for future in active_futures:
                    try:
                        task = active_futures[future]
                        completed_task = future.result(timeout=30)  # 给每个任务30秒的额外等待时间
                        
                        if completed_task.status == "已完成":
                            completed_count += 1
                            print(f"✓ 最终任务完成: {completed_task.filename}")
                            self.save_translation_result(completed_task, output_folder)
                        elif completed_task.status == "失败":
                            failed_count += 1
                            print(f"✗ 最终任务失败: {completed_task.filename} - {completed_task.error_message}")
                            
                    except concurrent.futures.TimeoutError:
                        failed_count += 1
                        task.status = "失败"
                        task.error_message = "任务最终等待超时"
                        print(f"✗ 最终任务超时: {task.filename}")
                        
                    except Exception as e:
                        failed_count += 1
                        task.status = "失败"
                        task.error_message = str(e)
                        print(f"✗ 最终任务处理出错: {task.filename} - {str(e)}")
            
            self.is_translating = False
            self.is_paused = False
            
            # 更新按钮状态
            self.root.after(0, lambda: self.start_btn.config(state='normal'))
            self.root.after(0, lambda: self.pause_btn.config(state='disabled', text="⏸️ 暂停翻译"))
            self.root.after(0, lambda: self.stop_btn.config(state='disabled'))
            
            # 显示完成统计
            print(f"\n{'#'*60}")
            print(f"# 批量翻译完成")
            print(f"# 成功: {completed_count} 个")
            print(f"# 失败: {failed_count} 个")
            print(f"# 总计: {len(self.translation_tasks)} 个")
            print(f"{'#'*60}\n")
            
            if failed_count > 0:
                self.root.after(0, lambda: self.file_status_var.set(f"⚠️ 翻译完成，{completed_count}个成功，{failed_count}个失败"))
            else:
                self.root.after(0, lambda: self.file_status_var.set(f"✅ 翻译完成！共{completed_count}个文件"))
            
        except Exception as e:
            self.is_translating = False
            self.is_paused = False
            
            # 更新按钮状态
            self.root.after(0, lambda: self.start_btn.config(state='normal'))
            self.root.after(0, lambda: self.pause_btn.config(state='disabled', text="⏸️ 暂停翻译"))
            self.root.after(0, lambda: self.stop_btn.config(state='disabled'))
            
            error_msg = f"❌ 翻译失败: {e}"
            print(f"\n{'#'*60}")
            print(f"# 批量翻译失败")
            print(f"# 错误: {str(e)}")
            print(f"{'#'*60}\n")
            
            import traceback
            traceback.print_exc()
            
            self.root.after(0, lambda: self.file_status_var.set(error_msg))
    
    def translate_single_file(self, task, config):
        """翻译单个文件 - 修复超时问题"""
        try:
            # 检查是否停止
            if not self.is_translating:
                task.status = "已取消"
                return task
            
            task.status = "翻译中"
            task.start_time = time.time()
            task.progress = 10
            
            print(f"\n{'='*60}")
            print(f"开始翻译文件: {task.filename}")
            print(f"文件大小: {len(task.content)} 字符")
            print(f"{'='*60}\n")
            
            # 预处理内容：检查是否需要添加标题
            processed_content = self.preprocess_content_with_title(task.content, task.filename)
            print(f"[预处理] 内容处理完成，最终大小: {len(processed_content)} 字符")
            
            # 加载模型
            print(f"[1/4] 加载模型: {config['model']} (端点: {config['endpoint']})")
            model_load(
                config['endpoint'],
                config['base_url'],
                config['model'],
                config['api_key'],
                config['temperature'],
                config['rpm']
            )
            task.progress = 20
            print(f"✓ 模型加载完成")
            
            # 定义进度回调函数
            def progress_callback(current, total, desc=""):
                """进度回调函数"""
                if not self.is_translating:
                    return
                
                # 计算进度百分比 (20-100)
                # 20% 已经用于模型加载
                # 剩余80%分配给翻译的3个阶段
                progress_percent = 20 + int((current / total) * 80)
                task.progress = progress_percent
                
                # 打印进度信息
                if desc:
                    print(f"[{current}/{total}] {desc} - 进度: {progress_percent}%")
            
            # 临时替换全局progress函数
            import sys
            # 获取process模块（已在文件顶部导入）
            process_module = sys.modules.get('process')
            if not process_module:
                # 如果没有找到，尝试从app目录导入
                import importlib.util
                spec = importlib.util.spec_from_file_location("process", os.path.join(os.path.dirname(__file__), 'app', 'process.py'))
                process_module = importlib.util.module_from_spec(spec)
                spec.loader.exec_module(process_module)
            
            original_progress = getattr(process_module, 'progress', None)
            
            # 创建新的progress函数
            def new_progress(progress_tuple, desc=""):
                if isinstance(progress_tuple, tuple) and len(progress_tuple) == 2:
                    current, total = progress_tuple
                    progress_callback(current, total, desc)
                # 如果有原始progress函数，也调用它
                if original_progress and callable(original_progress):
                    try:
                        original_progress(progress_tuple, desc=desc)
                    except:
                        pass
            
            # 替换progress函数
            process_module.progress = new_progress
            
            try:
                # 执行翻译（使用预处理后的内容）
                print(f"\n[2/4] 开始翻译流程...")
                if config['use_extra_endpoint']:
                    print(f"使用额外端点: {config['endpoint2']} / {config['model2']}")
                    init_translation, reflect_translation, final_translation = translator_sec(
                        endpoint2=config['endpoint2'],
                        base2=config['base_url2'],
                        model2=config['model2'],
                        api_key2=config['api_key2'],
                        source_lang=config['source_lang'],
                        target_lang=config['target_lang'],
                        source_text=processed_content,  # 使用预处理后的内容
                        country=config['country'],
                        max_tokens=config['max_tokens']
                    )
                else:
                    print(f"使用单一端点翻译")
                    init_translation, reflect_translation, final_translation = translator(
                        source_lang=config['source_lang'],
                        target_lang=config['target_lang'],
                        source_text=processed_content,  # 使用预处理后的内容
                        country=config['country'],
                        max_tokens=config['max_tokens']
                    )
                print(f"✓ 翻译流程完成")
            finally:
                # 恢复原始progress函数
                if original_progress:
                    process_module.progress = original_progress
            
            task.init_translation = init_translation
            task.reflect_translation = reflect_translation
            task.final_translation = final_translation
            task.progress = 100
            task.status = "已完成"
            task.end_time = time.time()
            
            elapsed = task.end_time - task.start_time
            print(f"\n{'='*60}")
            print(f"✅ 翻译完成: {task.filename}")
            print(f"耗时: {elapsed:.2f} 秒")
            print(f"初始翻译: {len(init_translation)} 字符")
            print(f"最终翻译: {len(final_translation)} 字符")
            print(f"{'='*60}\n")
            
        except Exception as e:
            task.status = "失败"
            task.error_message = str(e)
            task.progress = 0
            task.end_time = time.time()
            
            # 打印详细错误信息
            print(f"\n{'='*60}")
            print(f"❌ 翻译失败: {task.filename}")
            print(f"错误类型: {type(e).__name__}")
            print(f"错误信息: {str(e)}")
            print(f"{'='*60}\n")
            
            # 打印完整的错误堆栈
            import traceback
            traceback.print_exc()
        
        return task
    
    def preprocess_content_with_title(self, content, filename):
        """预处理内容：如果没有标题则添加文件名作为标题"""
        try:
            lines = content.strip().split('\n')
            if not lines:
                # 如果内容为空，只添加文件名作为标题
                print(f"[预处理] 内容为空，添加文件名作为标题: {filename}")
                return f"{filename}\n\n"
            
            first_line = lines[0].strip()
            
            # 检查第一行是否已经是标题
            has_title = (
                first_line and
                len(first_line) < 100 and
                (
                    'chapter' in first_line.lower() or
                    'prologue' in first_line.lower() or
                    'epilogue' in first_line.lower() or
                    'part' in first_line.lower() or
                    '章' in first_line or
                    filename.replace('.txt', '').replace('.md', '') in first_line  # 检查是否包含文件名
                )
            )
            
            if has_title:
                print(f"[预处理] 内容已有标题: {first_line}")
                return content
            else:
                # 没有标题，在开头添加文件名作为标题
                print(f"[预处理] 内容无标题，添加文件名: {filename}")
                processed_content = f"{filename}\n\n{content}"
                return processed_content
                
        except Exception as e:
            print(f"[预处理] 处理失败: {e}，使用原内容")
            return content
    
    def save_translation_result(self, task, output_folder):
        """保存翻译结果"""
        try:
            output_format = self.output_format_var.get()
            
            print(f"[3/4] 保存翻译结果: {task.filename}.{output_format}")
            
            if output_format == "txt":
                self.save_as_txt(task, output_folder)
            elif output_format == "docx":
                self.save_as_docx(task, output_folder)
            
            print(f"✓ 文件保存成功")
            
        except Exception as e:
            print(f"✗ 保存文件时出错: {e}")
            import traceback
            traceback.print_exc()
            
            task.status = "保存失败"
            task.error_message = f"保存文件时出错: {e}"
    
    def save_as_txt(self, task, output_folder):
        """保存为TXT格式 - 只输出最终翻译结果"""
        # 智能选择文件名：优先使用翻译内容中的章节名，否则使用文件名翻译
        output_filename = self.get_smart_filename(task)
        output_path = os.path.join(output_folder, output_filename)
        
        # 只保存最终翻译结果
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(task.final_translation)
        
        print(f"TXT文件已保存到: {output_path}")
    
    def get_smart_filename(self, task):
        """智能获取文件名：优先使用内容中的章节名，否则使用文件名翻译"""
        try:
            # 1. 尝试从翻译内容中提取章节名
            if task.final_translation:
                lines = task.final_translation.strip().split('\n')
                if lines:
                    first_line = lines[0].strip()
                    
                    # 检查第一行是否看起来像章节标题
                    # 章节标题特征：包含Chapter、第X章、较短（<100字符）
                    is_chapter_title = (
                        first_line and
                        len(first_line) < 100 and
                        (
                            'chapter' in first_line.lower() or
                            'prologue' in first_line.lower() or
                            'epilogue' in first_line.lower() or
                            'part' in first_line.lower() or
                            '章' in first_line
                        )
                    )
                    
                    if is_chapter_title:
                        print(f"[文件命名] 使用内容中的章节名: {first_line}")
                        
                        # 清理章节名，移除不适合文件名的字符
                        title = first_line.replace(':', ' -').replace('/', '-').replace('\\', '-')
                        title = title.replace('?', '').replace('*', '').replace('"', '').replace('<', '').replace('>', '').replace('|', '')
                        title = title.strip()
                        
                        # 限制文件名长度
                        if len(title) > 100:
                            title = title[:100]
                        
                        return f"{title}.txt"
            
            # 2. 如果内容中没有章节名，翻译原始文件名
            print(f"[文件命名] 内容中无章节名，翻译文件名: {task.filename}")
            return self.translate_filename_to_english(task.filename)
            
        except Exception as e:
            print(f"[文件命名] 获取文件名失败: {e}，使用文件名翻译")
            return self.translate_filename_to_english(task.filename)
    
    def translate_filename_to_english(self, chinese_filename):
        """将中文文件名翻译成英文"""
        try:
            # 检查文件名是否包含中文
            if not any('\u4e00' <= char <= '\u9fff' for char in chinese_filename):
                # 如果没有中文，直接返回
                return f"{chinese_filename}.txt"
            
            print(f"[文件名翻译] 原始文件名: {chinese_filename}")
            
            # 使用当前配置的API翻译文件名
            # 构建翻译提示
            prompt = f"请将以下中文标题翻译成英文，保持简洁专业的风格。只返回翻译结果，不要添加任何解释或标点符号：\n\n{chinese_filename}"
            
            # 调用translator进行翻译（使用简单的单次翻译）
            try:
                # 使用one_chunk_initial_translation进行快速翻译
                from app.process import one_chunk_initial_translation
                
                english_title = one_chunk_initial_translation(
                    source_lang="Chinese",
                    target_lang="English",
                    source_text=prompt
                ).strip()
                
                # 清理翻译结果
                # 移除可能的引号、句号等
                english_title = english_title.strip('"\'.,;:!?。，；：！？')
                
                # 清理不适合文件名的字符
                english_title = english_title.replace(':', ' -').replace('/', '-').replace('\\', '-')
                english_title = english_title.replace('?', '').replace('*', '').replace('"', '').replace('<', '').replace('>', '').replace('|', '')
                
                # 限制文件名长度
                if len(english_title) > 100:
                    english_title = english_title[:100]
                
                print(f"[文件名翻译] 英文文件名: {english_title}")
                
                return f"{english_title}.txt"
                
            except Exception as e:
                print(f"[文件名翻译] 翻译失败: {e}，使用原文件名")
                # 如果翻译失败，使用原文件名
                safe_name = chinese_filename.replace(':', ' -').replace('/', '-').replace('\\', '-')
                safe_name = safe_name.replace('?', '').replace('*', '').replace('"', '').replace('<', '').replace('>', '').replace('|', '')
                return f"{safe_name}.txt"
                
        except Exception as e:
            print(f"[文件名翻译] 处理失败: {e}")
            return f"{chinese_filename}.txt"
    
    def extract_title_from_translation(self, translation_text, fallback_name):
        """从翻译文本中提取标题作为文件名（已弃用，保留用于兼容）"""
        if not translation_text:
            return f"{fallback_name}.txt"
        
        # 尝试提取第一行作为标题
        lines = translation_text.strip().split('\n')
        if lines:
            first_line = lines[0].strip()
            
            # 如果第一行看起来像标题（包含Chapter、第X章等）
            if first_line and (
                'chapter' in first_line.lower() or 
                '章' in first_line or
                'part' in first_line.lower() or
                len(first_line) < 100  # 标题通常较短
            ):
                # 清理标题，移除不适合文件名的字符
                title = first_line.replace(':', ' -').replace('/', '-').replace('\\', '-')
                title = title.replace('?', '').replace('*', '').replace('"', '').replace('<', '').replace('>', '').replace('|', '')
                
                # 限制文件名长度
                if len(title) > 100:
                    title = title[:100]
                
                return f"{title}.txt"
        
        # 如果无法提取标题，使用原文件名
        return f"{fallback_name}.txt"
    
    def save_as_docx(self, task, output_folder):
        """保存为DOCX格式 - 只输出最终翻译结果"""
        try:
            from docx import Document
        except ImportError:
            # 如果没有python-docx，回退到txt格式
            print("警告: 缺少python-docx库，回退到TXT格式保存")
            self.save_as_txt(task, output_folder)
            return
        
        # 使用智能文件名（优先内容章节名，否则文件名翻译）
        output_filename = self.get_smart_filename(task).replace('.txt', '.docx')
        output_path = os.path.join(output_folder, output_filename)
        
        # 创建Word文档
        doc = Document()
        
        # 直接添加翻译内容，不添加额外的标题和信息
        # 按段落分割并添加
        paragraphs = task.final_translation.split('\n')
        for para_text in paragraphs:
            if para_text.strip():  # 跳过空行
                doc.add_paragraph(para_text)
            else:
                doc.add_paragraph()  # 保留空行
        
        # 保存文档
        doc.save(output_path)
        print(f"DOCX文件已保存到: {output_path}")
    
    def clear_tasks(self):
        """清空任务"""
        if self.is_translating:
            if messagebox.askyesno("确认", "翻译正在进行中，确定要清空任务吗？"):
                self.is_translating = False
                self.is_paused = False
                
                # 更新按钮状态
                self.start_btn.config(state='normal')
                self.pause_btn.config(state='disabled', text="⏸️ 暂停翻译")
                self.stop_btn.config(state='disabled')
            else:
                return
        
        self.translation_tasks.clear()
        self.task_counter = 0
        self.file_status_var.set("任务已清空")
        self.refresh_progress()
    
    def update_progress_display(self):
        """更新进度显示"""
        if not self.translation_tasks:
            self.overall_progress_var.set("暂无翻译任务")
            self.overall_progress_bar['value'] = 0
        else:
            completed = sum(1 for task in self.translation_tasks.values() if task.status == "已完成")
            total = len(self.translation_tasks)
            progress = (completed / total) * 100 if total > 0 else 0
            
            self.overall_progress_var.set(f"总体进度: {completed}/{total} 已完成 ({progress:.1f}%)")
            self.overall_progress_bar['value'] = progress
        
        # 更新详细进度
        self.refresh_progress()
        
        # 更新实时统计
        self.update_realtime_stats()
        
        # 每2秒更新一次
        self.root.after(2000, self.update_progress_display)
    
    def refresh_progress(self):
        """刷新进度显示"""
        # 清空现有项目
        for item in self.progress_tree.get_children():
            self.progress_tree.delete(item)
        
        # 添加任务进度
        for task in self.translation_tasks.values():
            elapsed_time = ""
            if task.start_time:
                if task.end_time:
                    elapsed = task.end_time - task.start_time
                    elapsed_time = f"{elapsed:.1f}s"
                else:
                    elapsed = time.time() - task.start_time
                    elapsed_time = f"{elapsed:.1f}s"
            
            error_msg = task.error_message[:50] + "..." if len(task.error_message) > 50 else task.error_message
            
            self.progress_tree.insert('', 'end', values=(
                task.filename,
                task.status,
                f"{task.progress}%",
                elapsed_time,
                error_msg
            ))
    
    def open_output_folder(self):
        """打开输出文件夹"""
        output_folder = self.output_folder_var.get()
        if output_folder and os.path.exists(output_folder):
            if sys.platform == "win32":
                os.startfile(output_folder)
            elif sys.platform == "darwin":
                os.system(f"open '{output_folder}'")
            else:
                os.system(f"xdg-open '{output_folder}'")
        else:
            messagebox.showerror("错误", "输出文件夹不存在")
    
    def save_config(self):
        """保存配置"""
        try:
            config = {
                'endpoint': self.endpoint_var.get(),
                'model': self.model_var.get(),
                'api_key': self.api_key_var.get(),
                'base_url': self.base_url_var.get(),
                'use_extra_endpoint': self.use_extra_endpoint_var.get(),
                'endpoint2': self.endpoint2_var.get(),
                'model2': self.model2_var.get(),
                'api_key2': self.api_key2_var.get(),
                'base_url2': self.base_url2_var.get(),
                'source_lang': self.source_lang_var.get(),
                'target_lang': self.target_lang_var.get(),
                'country': self.country_var.get(),
                'max_tokens': self.max_tokens_var.get(),
                'temperature': self.temperature_var.get(),
                'rpm': self.rpm_var.get(),
                'input_folder': self.input_folder_var.get(),
                'output_folder': self.output_folder_var.get(),
                'output_format': self.output_format_var.get(),
                'concurrent_tasks': self.concurrent_var.get(),
                'file_types': {k: v.get() for k, v in self.file_types.items()},
                # 新增性能优化设置
                'api_timeout': getattr(self, 'api_timeout_var', tk.IntVar(value=300)).get(),
                'performance_mode': getattr(self, 'performance_mode_var', tk.StringVar(value="平衡")).get(),
                'retry_count': getattr(self, 'retry_count_var', tk.IntVar(value=2)).get()
            }
            
            with open(CONFIG_FILE, 'w', encoding='utf-8') as f:
                json.dump(config, f, indent=2, ensure_ascii=False)
            
            self.api_status_var.set("✅ 配置保存成功")
            
        except Exception as e:
            self.api_status_var.set(f"❌ 保存失败: {e}")
    
    def load_config(self):
        """加载配置"""
        # 设置标志，防止 on_endpoint_change 覆盖模型名
        self.is_loading_config = True
        
        try:
            if os.path.exists(CONFIG_FILE):
                with open(CONFIG_FILE, 'r', encoding='utf-8') as f:
                    config = json.load(f)
                
                # 加载API配置
                self.endpoint_var.set(config.get('endpoint', 'OpenAI'))
                self.model_var.set(config.get('model', 'gpt-4o'))
                self.api_key_var.set(config.get('api_key', ''))
                self.base_url_var.set(config.get('base_url', ''))
                self.use_extra_endpoint_var.set(config.get('use_extra_endpoint', False))
                self.endpoint2_var.set(config.get('endpoint2', 'OpenAI'))
                self.model2_var.set(config.get('model2', 'gpt-4o'))
                self.api_key2_var.set(config.get('api_key2', ''))
                self.base_url2_var.set(config.get('base_url2', ''))
                
                # 加载翻译参数
                self.source_lang_var.set(config.get('source_lang', 'Chinese'))
                self.target_lang_var.set(config.get('target_lang', 'English'))
                self.country_var.set(config.get('country', 'United States'))
                self.max_tokens_var.set(config.get('max_tokens', 1000))
                self.temperature_var.set(config.get('temperature', 0.3))
                self.rpm_var.set(config.get('rpm', 60))
                
                # 加载文件设置
                self.input_folder_var.set(config.get('input_folder', ''))
                self.output_folder_var.set(config.get('output_folder', str(Path.home() / "Desktop" / "translations")))
                self.output_format_var.set(config.get('output_format', 'txt'))
                self.concurrent_var.set(config.get('concurrent_tasks', 5))
                
                # 加载文件类型
                file_types_config = config.get('file_types', {})
                for file_type, var in self.file_types.items():
                    var.set(file_types_config.get(file_type, file_type in ['txt', 'md', 'pdf', 'docx']))
                
                # 加载性能优化设置
                if hasattr(self, 'api_timeout_var'):
                    self.api_timeout_var.set(config.get('api_timeout', 300))  # 默认5分钟
                if hasattr(self, 'performance_mode_var'):
                    self.performance_mode_var.set(config.get('performance_mode', '平衡'))
                if hasattr(self, 'retry_count_var'):
                    self.retry_count_var.set(config.get('retry_count', 2))
                
                # 更新界面（显示/隐藏base_url字段）
                self.on_endpoint_change()
                self.on_endpoint2_change()
                self.toggle_extra_endpoint()
                
                self.api_status_var.set("✅ 配置加载成功")
                
        except Exception as e:
            self.api_status_var.set(f"⚠️ 配置加载失败: {e}")
        finally:
            # 加载完成后重置标志
            self.is_loading_config = False
    
    def on_closing(self):
        """关闭程序时的处理"""
        if self.is_translating:
            if messagebox.askyesno("确认退出", "翻译正在进行中，确定要退出吗？"):
                self.save_config()
                self.root.destroy()
        else:
            self.save_config()
            self.root.destroy()


def main():
    """主函数"""
    # 检查依赖
    try:
        import tkinter as tk
        from tkinter import ttk
    except ImportError:
        print("❌ 缺少 tkinter 库，请安装 Python 的 tkinter 支持")
        return
    
    # 创建主窗口
    root = tk.Tk()
    
    # 创建应用
    app = TranslationAgentGUI(root)
    
    # 运行应用
    root.mainloop()


if __name__ == "__main__":
    main()